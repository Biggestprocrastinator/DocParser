import base64
import os
import cv2
import numpy as np
import docx

def save_base64_to_tempfile(base64_string: str, temp_dir: str, file_name: str) -> str:
    """
    Decodes a base64 string and saves it to a temporary directory.
    Returns the absolute path to the saved file.
    """
    file_path = os.path.join(temp_dir, file_name)
    
    
    if "," in base64_string:
        base64_string = base64_string.split(",")[1]
        
    file_data = base64.b64decode(base64_string)
    
    with open(file_path, "wb") as f:
        f.write(file_data)
        
    return file_path

def clean_text(text: str) -> str:
    
    if not text:
        return ""
    
    # Strip common OCR table artifacts
    text = text.replace("|", " ").replace("_", " ")

    # Remove excessive newlines
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    cleaned = " ".join(lines)
    
    # Remove non-printable characters 
    cleaned = "".join(char for char in cleaned if char.isprintable())
    
    return cleaned

def preprocess_image(image_data: np.ndarray) -> np.ndarray:

    #Applies global pre-processing: Grayscale -> Gaussian Blur -> Otsu's Thresholding.
    
    if image_data is None or image_data.size == 0:
        return image_data

    # 1. Resize Image x2 using Cubic Interpolation to help tesserract
    resized = cv2.resize(image_data, None, fx=2.0, fy=2.0, interpolation=cv2.INTER_CUBIC)

    # 2. Convert to Grayscale
    if len(resized.shape) == 3 and resized.shape[2] == 3:
        gray = cv2.cvtColor(resized, cv2.COLOR_BGR2GRAY)
    elif len(resized.shape) == 3 and resized.shape[2] == 4:
        gray = cv2.cvtColor(resized, cv2.COLOR_BGRA2GRAY)
    else:
        gray = resized

    # 3. Apply Gaussian Blur to reduce noise
    blurred = cv2.GaussianBlur(gray, (5, 5), 0)

    # 4. Apply Otsu's Thresholding

    _, thresh = cv2.threshold(blurred, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)

    return thresh

def extract_text_from_docx(file_path: str) -> str:

    #Extracts text from a DOCX file directly from its XML structure.
    
    try:
        doc = docx.Document(file_path)
        full_text = []

        # 1. Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())

        # 2. Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    if cell.text.strip():
                        
                        row_data.append(cell.text.strip().replace('\n', ' ')) 
                if row_data:
                    full_text.append(" | ".join(row_data))

        return "\n".join(full_text)
    except Exception as e:
        print(f"DOCX Extraction Error: {e}")
        return ""
